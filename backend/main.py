"""
FastAPI application with all endpoints for the Legal AI system
"""
import os
import uuid
import shutil
from datetime import datetime
from typing import List, Optional, Dict, Any
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
import uvicorn

# Import models and schemas
from models.database import get_db, create_tables, Document, Session as DBSession
from models.schemas import (
    DocumentUploadRequest, SummarizeRequest, ExtractClausesRequest,
    DefineTermsRequest, AskQuestionRequest, RouteRequest,
    DocumentInfo, SummaryResponse, ClausesResponse, DefinitionsResponse,
    QAResponse, SessionState, RouteResponse, ErrorResponse
)

# Import agents
from agents.base_agent import AgentInput, AgentState
from agents.routing_state_agent import RoutingStateAgent
from agents.document_retrieval_agent import DocumentRetrievalAgent
from agents.summarization_agent import SummarizationAgent
from agents.clauses_extraction_agent import ClausesExtractionAgent
from agents.legal_definitions_agent import LegalDefinitionsAgent
from agents.qa_agent import QAAgent

from config import settings

# Create FastAPI app
app = FastAPI(
    title="Legal AI Document Analysis API",
    description="AI-powered legal document analysis and Q&A system",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify allowed origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize agents
routing_agent = RoutingStateAgent()
document_agent = DocumentRetrievalAgent()
summarization_agent = SummarizationAgent()
clauses_agent = ClausesExtractionAgent()
definitions_agent = LegalDefinitionsAgent()
qa_agent = QAAgent()

# Create tables on startup
create_tables()

# Ensure upload directory exists
os.makedirs(settings.upload_dir, exist_ok=True)


@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "message": "Legal AI Document Analysis API",
        "version": "1.0.0",
        "status": "active"
    }


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0"
    }


@app.post("/upload_document", response_model=DocumentInfo)
async def upload_document(
    file: UploadFile = File(...),
    user_id: Optional[str] = None,
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db)
):
    """Upload and process a legal document"""
    try:
        # Validate file type
        allowed_extensions = ['.pdf', '.docx', '.txt']
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Check file size
        if file.size > settings.max_file_size:
            raise HTTPException(
                status_code=400,
                detail=f"File too large. Maximum size: {settings.max_file_size / 1024 / 1024:.1f}MB"
            )
        
        # Generate document ID and save file
        document_id = str(uuid.uuid4())
        file_path = os.path.join(settings.upload_dir, f"{document_id}{file_ext}")
        
        # Save uploaded file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Create database record
        doc_record = Document(
            id=document_id,
            filename=file.filename,
            file_type=file_ext[1:],  # Remove the dot
            file_size=file.size,
            user_id=user_id,
            processing_status="pending"
        )
        
        db.add(doc_record)
        db.commit()
        db.refresh(doc_record)
        
        # Process document in background
        background_tasks.add_task(
            process_document_background,
            document_id=document_id,
            file_path=file_path,
            file_type=file_ext[1:]
        )
        
        return DocumentInfo(
            id=document_id,
            filename=file.filename,
            file_type=file_ext[1:],
            file_size=file.size,
            upload_timestamp=doc_record.upload_timestamp,
            processing_status="processing",
            metadata={}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


async def process_document_background(document_id: str, file_path: str, file_type: str):
    """Background task to process uploaded document"""
    try:
        print(f"Starting document processing for {document_id}, file: {file_path}")
        
        # Simple file processing with text extraction
        import os
        
        # Check if file exists
        if not os.path.exists(file_path):
            raise Exception(f"File not found: {file_path}")
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Extract text content based on file type
        full_text = ""
        content_preview = ""
        
        if file_type.lower() == 'txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    full_text = f.read()
                    content_preview = full_text[:500] + "..." if len(full_text) > 500 else full_text
            except:
                content_preview = "Text file uploaded successfully"
                full_text = "Text file content could not be extracted."
                
        elif file_type.lower() == 'pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    full_text = ""
                    for page in pdf_reader.pages:
                        full_text += page.extract_text() + "\n"
                    content_preview = full_text[:500] + "..." if len(full_text) > 500 else full_text
            except Exception as e:
                print(f"PDF extraction error: {e}")
                content_preview = "PDF document uploaded successfully"
                full_text = "PDF content extraction failed."
                
        elif file_type.lower() == 'docx':
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                full_text = ""
                for paragraph in doc.paragraphs:
                    full_text += paragraph.text + "\n"
                content_preview = full_text[:500] + "..." if len(full_text) > 500 else full_text
            except Exception as e:
                print(f"DOCX extraction error: {e}")
                content_preview = "Word document uploaded successfully"
                full_text = "DOCX content extraction failed."
        else:
            full_text = f"Unsupported file type: {file_type}"
            content_preview = f"File of type {file_type} uploaded successfully"
        
        # Store the full text in a simple way (we'll save it to a text file for now)
        text_storage_path = file_path + ".extracted.txt"
        try:
            with open(text_storage_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            print(f"Extracted text saved to: {text_storage_path}")
        except Exception as e:
            print(f"Failed to save extracted text: {e}")
        
        # Update database record with success
        db = next(get_db())
        doc_record = db.query(Document).filter(Document.id == document_id).first()
        
        if doc_record:
            doc_record.processing_status = "completed"
            doc_record.document_metadata = {
                'file_size_bytes': file_size,
                'processed_at': datetime.now().isoformat(),
                'status': 'basic_processing_complete',
                'text_length': len(full_text),
                'text_storage_path': text_storage_path
            }
            doc_record.content_preview = content_preview
            print(f"Document {document_id} processed successfully")
            db.commit()
        
        db.close()
        
    except Exception as e:
        print(f"Exception in process_document_background for {document_id}: {str(e)}")
        print(f"Exception type: {type(e)}")
        import traceback
        traceback.print_exc()
        
        # Update database with error status
        try:
            db = next(get_db())
            doc_record = db.query(Document).filter(Document.id == document_id).first()
            if doc_record:
                doc_record.processing_status = "failed"
                doc_record.document_metadata = {'error': str(e)}
                db.commit()
            db.close()
        except Exception as db_error:
            print(f"Database error: {str(db_error)}")


@app.get("/document/{document_id}", response_model=DocumentInfo)
async def get_document_info(document_id: str, db: Session = Depends(get_db)):
    """Get document information"""
    doc_record = db.query(Document).filter(Document.id == document_id).first()
    
    if not doc_record:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return DocumentInfo(
        id=doc_record.id,
        filename=doc_record.filename,
        file_type=doc_record.file_type,
        file_size=doc_record.file_size,
        upload_timestamp=doc_record.upload_timestamp,
        processing_status=doc_record.processing_status,
        metadata=doc_record.document_metadata or {},
        content_preview=doc_record.content_preview
    )


@app.post("/summarize", response_model=SummaryResponse)
async def summarize_document(request: SummarizeRequest):
    """Generate document summary"""
    try:
        # Create agent state
        state = AgentState(session_id=str(uuid.uuid4()))
        
        # Prepare input
        input_data = AgentInput(
            query="summarize this document",
            document_id=request.document_id,
            parameters={
                'summary_type': request.summary_type,
                'section_name': request.section_name
            }
        )
        
        # Run summarization agent
        result = await summarization_agent.run(input_data, state)
        
        if not result.success:
            raise HTTPException(status_code=500, detail=result.error_message)
        
        return SummaryResponse(
            document_id=request.document_id,
            summary_type=request.summary_type,
            summary_text=result.result.get('summary_text', ''),
            section_name=request.section_name,
            key_points=result.result.get('key_points', []),
            created_at=datetime.now()
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Summarization failed: {str(e)}")


@app.post("/extract_clauses", response_model=ClausesResponse)
async def extract_clauses(request: ExtractClausesRequest):
    """Extract legal clauses from document"""
    try:
        # Create agent state
        state = AgentState(session_id=str(uuid.uuid4()))
        
        # Prepare input
        input_data = AgentInput(
            query="extract clauses from this document",
            document_id=request.document_id,
            parameters={
                'clause_types': [ct.value for ct in request.clause_types] if request.clause_types else None
            }
        )
        
        # Run clauses extraction agent
        result = await clauses_agent.run(input_data, state)
        
        if not result.success:
            raise HTTPException(status_code=500, detail=result.error_message)
        
        # Format clauses
        clauses = []
        extracted_clauses = result.result.get('extracted_clauses', {})
        
        for clause_type, clause_list in extracted_clauses.items():
            for clause in clause_list:
                clauses.append({
                    'id': str(uuid.uuid4()),
                    'clause_type': clause.get('clause_type', clause_type),
                    'clause_text': clause.get('text', ''),
                    'simplified_explanation': clause.get('description', ''),
                    'page_number': clause.get('page_number'),
                    'section': clause.get('section'),
                    'importance_score': clause.get('importance_score', 5)
                })
        
        return ClausesResponse(
            document_id=request.document_id,
            clauses=clauses,
            extraction_timestamp=datetime.now()
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Clause extraction failed: {str(e)}")


@app.post("/define_terms", response_model=DefinitionsResponse)
async def define_legal_terms(request: DefineTermsRequest):
    """Define legal terms"""
    try:
        # Create agent state
        state = AgentState(session_id=str(uuid.uuid4()))
        
        # Prepare input
        input_data = AgentInput(
            query=f"define these terms: {', '.join(request.terms)}",
            document_id=request.document_id,
            parameters={
                'terms': request.terms,
                'context': request.context
            }
        )
        
        # Run legal definitions agent
        result = await definitions_agent.run(input_data, state)
        
        if not result.success:
            raise HTTPException(status_code=500, detail=result.error_message)
        
        # Format definitions
        definitions = []
        for definition in result.result.get('definitions', []):
            definitions.append({
                'term': definition.get('term', ''),
                'legal_definition': definition.get('legal_definition', ''),
                'simple_definition': definition.get('simple_definition', ''),
                'examples': definition.get('examples', []),
                'related_terms': definition.get('related_terms', [])
            })
        
        return DefinitionsResponse(
            definitions=definitions,
            context_document_id=request.document_id
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Term definition failed: {str(e)}")


@app.post("/ask", response_model=QAResponse)
async def ask_question(request: AskQuestionRequest):
    """Ask a question about the document"""
    try:
        # Create agent state
        state = AgentState(session_id=str(uuid.uuid4()))
        
        # Prepare input
        input_data = AgentInput(
            query=request.question,
            document_id=request.document_id,
            parameters={
                'conversation_history': request.conversation_history
            }
        )
        
        # Run Q&A agent
        result = await qa_agent.run(input_data, state)
        
        if not result.success:
            raise HTTPException(status_code=500, detail=result.error_message)
        
        # Format citations
        citations = []
        for citation in result.result.get('citations', []):
            # Ensure relevance score is within valid range (0.0 to 1.0)
            relevance_score = citation.get('relevance_score', 0.0)
            relevance_score = max(0.0, min(1.0, relevance_score))  # Clamp to [0, 1]
            
            citations.append({
                'text': citation.get('text', ''),
                'page_number': citation.get('page_number'),
                'section': citation.get('section'),
                'relevance_score': relevance_score
            })
        
        return QAResponse(
            question=request.question,
            answer=result.result.get('answer', ''),
            citations=citations,
            confidence_score=result.result.get('confidence_score', 0.0),
            related_questions=result.result.get('related_questions', [])
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Question answering failed: {str(e)}")


@app.post("/route", response_model=RouteResponse)
async def route_request(request: RouteRequest):
    """Route request to appropriate agent"""
    try:
        # Create agent state
        state = AgentState(session_id=request.session_id)
        
        # Prepare input
        input_data = AgentInput(
            query=request.query,
            document_id=request.document_id,
            parameters={'action': 'route'}
        )
        
        # Run routing agent
        result = await routing_agent.run(input_data, state)
        
        if not result.success:
            raise HTTPException(status_code=500, detail=result.error_message)
        
        return RouteResponse(
            agent_used=result.result.get('agent_used', ''),
            result=result.result.get('result', {}),
            next_suggested_action=result.result.get('next_suggested_actions', [{}])[0].get('description') if result.result.get('next_suggested_actions') else None,
            confidence=result.result.get('confidence', 0.0)
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Request routing failed: {str(e)}")


@app.post("/session", response_model=SessionState)
async def create_session(user_id: Optional[str] = None):
    """Create a new session"""
    try:
        # Create agent state
        state = AgentState(session_id=str(uuid.uuid4()))
        
        # Prepare input
        input_data = AgentInput(
            query="create session",
            parameters={
                'action': 'create_session',
                'user_id': user_id
            }
        )
        
        # Run routing agent to create session
        result = await routing_agent.run(input_data, state)
        
        if not result.success:
            raise HTTPException(status_code=500, detail=result.error_message)
        
        session_data = result.result
        
        return SessionState(
            session_id=session_data['session_id'],
            user_id=user_id,
            created_at=datetime.fromisoformat(session_data['created_at']),
            updated_at=datetime.fromisoformat(session_data['created_at']),
            is_active=True,
            context={},
            uploaded_documents=[],
            conversation_count=0
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Session creation failed: {str(e)}")


@app.get("/session/{session_id}", response_model=SessionState)
async def get_session_state(session_id: str):
    """Get session state"""
    try:
        # Create agent state
        state = AgentState(session_id=session_id)
        
        # Prepare input
        input_data = AgentInput(
            query="get session state",
            parameters={
                'action': 'get_session',
                'session_id': session_id
            }
        )
        
        # Run routing agent to get session
        result = await routing_agent.run(input_data, state)
        
        if not result.success:
            raise HTTPException(status_code=404, detail="Session not found")
        
        session_data = result.result
        
        return SessionState(
            session_id=session_data['session_id'],
            user_id=session_data['user_id'],
            created_at=datetime.fromisoformat(session_data['created_at']),
            updated_at=datetime.fromisoformat(session_data['updated_at']),
            is_active=session_data['is_active'],
            context=session_data['context'],
            uploaded_documents=session_data['uploaded_documents'],
            conversation_count=session_data['conversation_count']
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get session state: {str(e)}")


@app.get("/documents")
async def list_documents(user_id: Optional[str] = None, db: Session = Depends(get_db)):
    """List all documents, optionally filtered by user"""
    try:
        query = db.query(Document)
        
        if user_id:
            query = query.filter(Document.user_id == user_id)
        
        documents = query.order_by(Document.upload_timestamp.desc()).all()
        
        return [
            {
                'id': doc.id,
                'filename': doc.filename,
                'file_type': doc.file_type,
                'file_size': doc.file_size,
                'upload_timestamp': doc.upload_timestamp.isoformat(),
                'processing_status': doc.processing_status,
                'user_id': doc.user_id
            }
            for doc in documents
        ]
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to list documents: {str(e)}")


@app.delete("/document/{document_id}")
async def delete_document(document_id: str, db: Session = Depends(get_db)):
    """Delete a document"""
    try:
        doc_record = db.query(Document).filter(Document.id == document_id).first()
        
        if not doc_record:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Delete file if it exists
        file_path = os.path.join(settings.upload_dir, f"{document_id}.{doc_record.file_type}")
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Delete database record
        db.delete(doc_record)
        db.commit()
        
        return {"message": "Document deleted successfully"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")


if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host=settings.api_host,
        port=settings.api_port,
        reload=settings.debug
    )